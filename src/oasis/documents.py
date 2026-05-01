from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def _ensure_docx(path: str) -> Path:
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        from pdf2docx import Converter
        out = p.with_suffix(".docx")
        cv = Converter(str(p))
        cv.convert(str(out), start=0, end=None)
        cv.close()
        return out
    return p


def extract_paragraphs(path: str) -> list[str]:
    """Return list of paragraph texts (index = paragraph label used by Claude)."""
    doc = Document(str(_ensure_docx(path)))
    return [para.text for para in doc.paragraphs]


def apply_changes(path: str, changes: list[dict], output_path: str) -> None:
    """
    Apply Claude's changes to the document, preserving all Run-level formatting.
    changes: [{"index": N, "new_text": "..."}]
    """
    src = _ensure_docx(path)
    doc = Document(str(src))
    paras = doc.paragraphs

    for change in changes:
        idx: int = change["index"]
        new_text: str = change["new_text"]
        if idx >= len(paras):
            continue
        para = paras[idx]
        runs = para.runs
        if not runs:
            # Paragraph has no runs; add one using paragraph's default style
            para.clear()
            para.add_run(new_text)
            continue
        # Put all new text into the first run, preserve its formatting, clear the rest
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""

    doc.save(output_path)
